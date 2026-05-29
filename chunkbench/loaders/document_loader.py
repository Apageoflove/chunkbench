from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

import markdown
from bs4 import BeautifulSoup


@dataclass
class Document:
    doc_id: str
    text: str
    source: str
    doc_type: str = "general"
    metadata: Dict = field(default_factory=dict)


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md", ".html"}

    def __init__(self, doc_type: str = "general"):
        self.doc_type = doc_type

    def load_file(self, path: str) -> Document:
        ext = Path(path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")

        text = self._extract_text(path, ext)
        doc_id = Path(path).stem

        return Document(
            doc_id=doc_id,
            text=text,
            source=str(path),
            doc_type=self.doc_type,
            metadata={"file_extension": ext},
        )

    def load_dir(self, dir_path: str) -> List[Document]:
        documents = []
        for entry in sorted(Path(dir_path).iterdir()):
            if entry.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                documents.append(self.load_file(str(entry)))
        return documents

    def _extract_text(self, path: str, ext: str) -> str:
        if ext == ".txt":
            return Path(path).read_text(encoding="utf-8")

        if ext == ".md":
            md_text = Path(path).read_text(encoding="utf-8")
            html = markdown.markdown(md_text)
            return BeautifulSoup(html, "html.parser").get_text(separator="\n")

        if ext == ".html":
            html_text = Path(path).read_text(encoding="utf-8")
            return BeautifulSoup(html_text, "html.parser").get_text(separator="\n")

        if ext == ".pdf":
            return self._extract_pdf(path)

        if ext == ".docx":
            return self._extract_docx(path)

        raise ValueError(f"No extractor for {ext}")

    @staticmethod
    def _extract_pdf(path: str) -> str:
        from pypdf import PdfReader

        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(path: str) -> str:
        import docx

        doc = docx.Document(path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
